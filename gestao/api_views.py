import os
from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes as perm_classes
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from django.db.models import Count, Sum, Q, F
from django.db.models.functions import TruncMonth, TruncYear, ExtractMonth, ExtractYear
from .models import Ocorrencia, Noticiado, Material, RegistroHistorico, LoteIncineracao, NaturezaPenal, EquipePM, CATEGORIA_CHOICES, DROGAS_CHOICES, STATUS_CUSTODIA_CHOICES, VARA_CHOICES
from .api_serializers import (
    OcorrenciaSerializer, NoticiadoSerializer, MaterialSerializer, LoteIncineracaoSerializer,
    NaturezaPenalSerializer, EquipePMSerializer
)
from .api_filters import MaterialFilter
from .documentos_services import (
    gerar_oficio_materiais_gerais, gerar_recibo_entrada_pdf, 
    gerar_capa_lote_pdf, gerar_relatorio_filtrado_pdf
)
from django.conf import settings

def get_auth_user(request):
    return request.user if request.user.is_authenticated else None

class OcorrenciaViewSet(viewsets.ModelViewSet):
    queryset = Ocorrencia.objects.all().order_by('-data_criacao')
    serializer_class = OcorrenciaSerializer
    permission_classes = [AllowAny]

    def create(self, request, *args, **kwargs):
        print(f"[DEBUG] POST data: {request.data}")
        try:
            serializer = self.get_serializer(data=request.data)
            if not serializer.is_valid():
                print(f"[DEBUG] Validation errors: {serializer.errors}")
            return super().create(request, *args, **kwargs)
        except Exception as e:
            print(f"[DEBUG] Exception: {e}")
            raise

    def perform_create(self, serializer):
        serializer.context['request'] = self.request
        serializer.save(criado_por=get_auth_user(self.request))

    @action(detail=True, methods=['get'])
    def imprimir_recibo(self, request, pk=None):
        ocorrencia = self.get_object()
        caminho_relativo = gerar_recibo_entrada_pdf(ocorrencia)
        file_url = request.build_absolute_uri(settings.MEDIA_URL + caminho_relativo)
        return Response({'url': file_url})

    @action(detail=False, methods=['get'])
    def buscar_por_bou(self, request):
        bou = request.query_params.get('bou')
        if not bou:
            return Response({"error": "BOU não informado."}, status=400)
        
        try:
            ocorrencia = Ocorrencia.objects.get(bou=bou)
            serializer = self.get_serializer(ocorrencia)
            return Response(serializer.data)
        except Ocorrencia.DoesNotExist:
            return Response({"error": "BOU não encontrado."}, status=404)

class NoticiadoViewSet(viewsets.ModelViewSet):
    queryset = Noticiado.objects.all()
    serializer_class = NoticiadoSerializer
    permission_classes = [AllowAny]

    def perform_create(self, serializer):
        serializer.save(criado_por=get_auth_user(self.request))

class MaterialViewSet(viewsets.ModelViewSet):
    queryset = Material.objects.all().order_by('-data_criacao')
    serializer_class = MaterialSerializer
    permission_classes = [AllowAny]
    filterset_class = MaterialFilter

    @action(detail=False, methods=['get'])
    def estatisticas(self, request):
        """
        Retorna estatísticas agregadas respeitando TODOS os filtros aplicados.
        Endpoint: GET /api/materiais/estatisticas/?categoria=ENTORPECENTE&ano=2026&...
        """
        # Aplica os mesmos filtros do FilterSet
        qs = self.filter_queryset(self.get_queryset())

        total = qs.count()
        
        # --- Por Categoria ---
        por_categoria = list(
            qs.values('categoria')
            .annotate(total=Count('id'))
            .order_by('-total')
        )
        # Enriquece com label display
        cat_map = dict(CATEGORIA_CHOICES)
        for item in por_categoria:
            item['label'] = cat_map.get(item['categoria'], item['categoria'])

        # --- Por Substância (apenas entorpecentes) ---
        por_substancia = list(
            qs.filter(categoria='ENTORPECENTE')
            .values('substancia')
            .annotate(total=Count('id'))
            .order_by('-total')
        )
        subst_map = dict(DROGAS_CHOICES)
        for item in por_substancia:
            item['label'] = subst_map.get(item['substancia'], item['substancia'] or 'Não especificada')

        # --- Por Status ---
        por_status = list(
            qs.values('status')
            .annotate(total=Count('id'))
            .order_by('-total')
        )
        status_map = dict(STATUS_CUSTODIA_CHOICES)
        for item in por_status:
            item['label'] = status_map.get(item['status'], item['status'])

        # --- Por Vara ---
        por_vara = list(
            qs.values('noticiado__ocorrencia__vara')
            .annotate(total=Count('id'))
            .order_by('-total')
        )
        vara_map = dict(VARA_CHOICES)
        for item in por_vara:
            vara_key = item['noticiado__ocorrencia__vara']
            item['vara'] = vara_key
            item['label'] = vara_map.get(vara_key, vara_key or 'Não definida')

        # --- Evolução Mensal (últimos 12 meses ou período filtrado) ---
        por_mes = list(
            qs.annotate(
                mes=TruncMonth('data_criacao')
            )
            .values('mes')
            .annotate(total=Count('id'))
            .order_by('mes')
        )
        for item in por_mes:
            if item['mes']:
                item['mes_label'] = item['mes'].strftime('%m/%Y')
                item['mes'] = item['mes'].isoformat()

        # --- Top Autores/Réus ---
        top_autores = list(
            qs.values('noticiado__nome')
            .annotate(total=Count('id'))
            .order_by('-total')[:15]
        )
        for item in top_autores:
            item['nome'] = item.pop('noticiado__nome', 'Desconhecido')

        # --- Por Unidade PM ---
        por_unidade = list(
            qs.values('noticiado__ocorrencia__unidade_origem')
            .annotate(total=Count('id'))
            .order_by('-total')
        )
        for item in por_unidade:
            item['unidade'] = item.pop('noticiado__ocorrencia__unidade_origem', 'N/I')

        # --- Resumo numérico ---
        entorpecentes_count = qs.filter(categoria='ENTORPECENTE').count()
        materiais_gerais_count = qs.filter(categoria__in=['SOM', 'FACA', 'SIMULACRO', 'OUTROS']).count()
        dinheiro_count = qs.filter(categoria='DINHEIRO').count()
        incinerados_count = qs.filter(status='INCINERADO').count()
        no_cofre_count = qs.filter(status__in=['ARMAZENADO', 'AUTORIZADO']).count()
        entregues_count = qs.filter(status='ENTREGUE_AO_JUDICIARIO').count()
        
        # BOU únicos
        bous_unicos = qs.values('noticiado__ocorrencia__bou').distinct().count()
        
        # Peso total (entorpecentes)
        from django.db.models import DecimalField
        from django.db.models.functions import Coalesce
        peso_total = qs.filter(categoria='ENTORPECENTE').aggregate(
            total=Sum(Coalesce('peso_real', 'peso_estimado', output_field=DecimalField()))
        )['total'] or 0

        # --- Opções de filtro disponíveis (para popular os selects do frontend) ---
        categorias_disponiveis = [{'value': k, 'label': v} for k, v in CATEGORIA_CHOICES]
        substancias_disponiveis = [{'value': k, 'label': v} for k, v in DROGAS_CHOICES]
        status_disponiveis = [{'value': k, 'label': v} for k, v in STATUS_CUSTODIA_CHOICES]
        varas_disponiveis = [{'value': k, 'label': v} for k, v in VARA_CHOICES]
        anos_disponiveis = list(
            Material.objects.annotate(ano=ExtractYear('data_criacao'))
            .values_list('ano', flat=True)
            .distinct()
            .order_by('-ano')
        )

        return Response({
            'total': total,
            'resumo': {
                'entorpecentes': entorpecentes_count,
                'materiais_gerais': materiais_gerais_count,
                'dinheiro': dinheiro_count,
                'incinerados': incinerados_count,
                'no_cofre': no_cofre_count,
                'entregues_judiciario': entregues_count,
                'bous_unicos': bous_unicos,
                'peso_total_gramas': float(peso_total),
            },
            'por_categoria': por_categoria,
            'por_substancia': por_substancia,
            'por_status': por_status,
            'por_vara': por_vara,
            'por_mes': por_mes,
            'top_autores': top_autores,
            'por_unidade': por_unidade,
            'opcoes_filtro': {
                'categorias': categorias_disponiveis,
                'substancias': substancias_disponiveis,
                'status': status_disponiveis,
                'varas': varas_disponiveis,
                'anos': anos_disponiveis,
            }
        })

    def perform_create(self, serializer):
        user = get_auth_user(self.request)
        material = serializer.save(criado_por=user)
        RegistroHistorico.objects.create(
            material=material,
            criado_por=user,
            status_na_epoca=material.status,
            observacao="Material recebido pelo cartório."
        )

    def perform_update(self, serializer):
        user = get_auth_user(self.request)
        material = serializer.instance
        status_antigo = material.status
        material = serializer.save(criado_por=user)
        
        if status_antigo != material.status:
            RegistroHistorico.objects.create(
                material=material,
                criado_por=user,
                status_na_epoca=material.status,
                observacao=f"Status alterado de {status_antigo} para {material.status}"
            )
            
            # Gatilho de Atribuição Automática de Lote
            if material.status == 'AUTORIZADO' and not material.lote:
                self.auto_atribuir_lote(material, user)

    def auto_atribuir_lote(self, material, user):
        """Atribui material a um lote aberto (máx 20 processos por lote)"""
        bou = material.noticiado.ocorrencia.bou
        
        # 1. Verificar se algum material deste BOU já está em um lote aberto
        material_mesmo_bou = Material.objects.filter(
            noticiado__ocorrencia__bou=bou, 
            lote__status='ABERTO'
        ).exclude(lote=None).first()
        
        if material_mesmo_bou:
            material.lote = material_mesmo_bou.lote
            material.status = 'AGUARDANDO_INCINERACAO'
            material.save()
            return
        
        # 2. Procurar o último lote aberto
        lote = LoteIncineracao.objects.filter(status='ABERTO').order_by('data_criacao').last()
        
        # 3. Verificar quantidade de processos no lote
        if lote:
            processos_no_lote = Material.objects.filter(lote=lote).values_list('noticiado__ocorrencia__bou', flat=True).distinct().count()
            if processos_no_lote >= 20:
                lote = None # Criar novo se houver 20 processos
        
        # 4. Criar novo lote se necessário
        if not lote:
            from datetime import datetime
            prefix = f"LOTE-{datetime.now().strftime('%Y%m')}"
            count = LoteIncineracao.objects.filter(identificador__startswith=prefix).count() + 1
            identificador = f"{prefix}-{count:03d}"
            lote = LoteIncineracao.objects.create(
                identificador=identificador,
                status='ABERTO',
                criado_por=user
            )

        material.lote = lote
        material.status = 'AGUARDANDO_INCINERACAO'
        material.save()

    @action(detail=True, methods=['get'])
    def gerar_oficio_word(self, request, pk=None):
        material = self.get_object()
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from django.utils import timezone
        except ImportError:
            return Response({"error": "Biblioteca python-docx não instalada."}, status=500)
        
        ocorrencia = material.noticiado.ocorrencia if material.noticiado else None
        
        doc = Document()
        
        # Cabeçalho
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("POLÍCIA MILITAR DO PARANÁ")
        run.bold = True
        run.font.size = Pt(14)
        
        subheader = doc.add_paragraph()
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subheader.add_run("6º BATALHÃO DE POLÍCIA MILITAR - CASCAVEL/PR")
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Título do Ofício
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("OFÍCIO DE REMESSA")
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        # Número do ofício
        num_oficio = doc.add_paragraph()
        run = num_oficio.add_run(f"Nº do Ofício: {material.eprotocolo_geral or '[A definir]'}")
        run.bold = True
        
        # Data
        data_oficio = doc.add_paragraph()
        data_oficio.add_run(f"Cascavel/PR, {timezone.now().strftime('%d de %B de %Y')}")
        
        doc.add_paragraph()
        
        # Destinatário
        dest = doc.add_paragraph()
        dest.add_run("À Excelentíssima Senhora Juíza de Direito da ")
        dest.add_run(f"{ocorrencia.get_vara_display() if ocorrencia and ocorrencia.vara else '___ª Vara Criminal'}")
        dest.add_run("\nJuízo da Vara de Execuções Penais e Cartório")
        dest.add_run("\nCascavel/PR")
        
        doc.add_paragraph()
        
        # Assunto
        assunt = doc.add_paragraph()
        run = assunt.add_run("ASSUNTO: ")
        run.bold = True
        assunt.add_run("Remessa de Materiais Apreendidos para Destruição")
        
        doc.add_paragraph()
        
        # Corpo do ofício
        corpo = doc.add_paragraph()
        corpo.add_run(" Senhora Juíza,")
        corpo.add_run("\n\n")
        corpo.add_run("          Encaminho a Vossa Excelência o presente ofício acompanhado do ")
        corpo.add_run(f"material apreendido e armazenado nesta Unidade, conforme descrito abaixo:")
        
        doc.add_paragraph()
        
        # Tabela de materiais
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['Lacre', 'Substância', 'Peso', 'BOU', 'Processo']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        row_cells = table.add_row().cells
        row_cells[0].text = material.numero_lacre or 'N/I'
        row_cells[1].text = material.substancia.replace('_', ' ') if material.substancia else material.get_categoria_display()
        row_cells[2].text = material.peso_formatado() if hasattr(material, 'peso_formatado') else f"{material.peso_estimado or 'N/I'} {material.unidade or ''}"
        row_cells[3].text = ocorrencia.bou if ocorrencia else 'N/I'
        row_cells[4].text = ocorrencia.processo if ocorrencia else 'N/I'
        
        doc.add_paragraph()
        
        # Descrição
        if material.descricao_geral:
            desc = doc.add_paragraph()
            run = desc.add_run("DESCRIÇÃO DO MATERIAL: ")
            run.bold = True
            desc.add_run(f"\n{material.descricao_geral}")
            doc.add_paragraph()
        
        # Encerramento
        enc = doc.add_paragraph()
        enc.add_run("          Solicitamos que, após análise e despacho judicial autorizando a ")
        enc.add_run("destruição dos materiais ora encaminhados, seja retornado o presente ofício ")
        enc.add_run("acompanhado do Termo de Destruição devidamente assinado, para que possamos ")
        enc.add_run("proceder com a incineração.")
        
        doc.add_paragraph()
        
        # Atenciosamente
        atc = doc.add_paragraph()
        atc.add_run("          Atenciosamente,")
        atc.add_run("\n\n\n\n")
        atc.add_run("_______________________________________")
        atc.add_run("\nResponsável pelo Cartório do 6º BPM")
        
        doc.add_paragraph()
        
        # Rodapé
        rodape = doc.add_paragraph()
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rodape.add_run("6º Batalhão de Polícia Militar - Cascavel/PR")
        run.font.size = Pt(10)
        run2 = rodape.add_run("\nRua Rio Grande do Sul, 3450 - Centro - CEP 85801-000")
        run2.font.size = Pt(9)
        
        # Salvar
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        filename = f"oficio_{material.id}_{timestamp}.docx"
        pasta = os.path.join(settings.MEDIA_ROOT, 'oficios')
        os.makedirs(pasta, exist_ok=True)
        filepath = os.path.join(pasta, filename)
        
        doc.save(filepath)
        
        file_url = request.build_absolute_uri(settings.MEDIA_URL + 'oficios/' + filename)
        return Response({'url': file_url, 'filename': filename})

    @action(detail=False, methods=['post'])
    def gerar_oficio(self, request):
        ids_materiais = request.data.get('materiais_ids', [])
        print(f"[DEBUG] IDs recebidos: {ids_materiais} (Tipo: {type(ids_materiais)})")
        
        materiais = Material.objects.filter(id__in=ids_materiais, categoria__in=['SOM', 'FACA', 'SIMULACRO', 'OUTROS'])
        print(f"[DEBUG] Materiais encontrados: {materiais.count()}")
        
        if not materiais.exists():
            return Response({"error": "Nenhum material geral selecionado válido."}, status=status.HTTP_400_BAD_REQUEST)
        
        # Gera o PDF via service
        pdf_path = gerar_oficio_materiais_gerais(materiais, get_auth_user(request))
        
        # Atualiza o status
        for material in materiais:
            material.status = 'OFICIO_GERADO'
            material.save()
            RegistroHistorico.objects.create(
                material=material, criado_por=get_auth_user(request), 
                status_na_epoca='OFICIO_GERADO', observacao="Ofício de remessa gerado ao judiciário."
            )

        file_url = request.build_absolute_uri(settings.MEDIA_URL + f"oficios/{pdf_path.split('/')[-1]}")
        return Response({"message": "Ofício gerado", "file_url": file_url})

    @action(detail=False, methods=['post'])
    def gerar_relatorio(self, request):
        """
        Gera relatório PDF filtrado por tipo (inventario, incineracao, custodia, remessa).
        Body: { tipo: 'inventario'|'incineracao'|'custodia'|'remessa', filtros: {...} }
        """
        tipo = request.data.get('tipo', 'inventario')
        filtros = request.data.get('filtros', {})
        
        # Aplica filtros
        qs = Material.objects.all()
        
        if filtros.get('ano'):
            qs = qs.filter(data_criacao__year=filtros['ano'])
        if filtros.get('categoria'):
            qs = qs.filter(categoria=filtros['categoria'])
        if filtros.get('status'):
            qs = qs.filter(status=filtros['status'])
            
        # Mapeia tipo para categoria/status de filtro
        if tipo == 'incineracao':
            qs = qs.filter(status__in=['AUTORIZADO', 'AGUARDANDO_INCINERACAO'])
        elif tipo == 'custodia':
            qs = qs.filter(status__in=['ARMAZENADO', 'AUTORIZADO'])
        elif tipo == 'remessa':
            qs = qs.filter(categoria__in=['SOM', 'FACA', 'SIMULACRO', 'OUTROS'])
            
        if not qs.exists():
            return Response({"error": "Nenhum material encontrado com os filtros selecionados."}, status=400)
            
        # Gera o PDF
        filtros_labels = {
            'tipo': tipo,
            'ano': filtros.get('ano', ''),
            'categoria': filtros.get('categoria', ''),
            'status': filtros.get('status', ''),
        }
        pdf_path = gerar_relatorio_filtrado_pdf(qs, filtros_labels, tipo)
        
        file_url = request.build_absolute_uri(settings.MEDIA_URL + pdf_path)
        return Response({"message": "Relatório gerado", "file_url": file_url})

    @action(detail=True, methods=['post'])
    def conferir_fisicamente(self, request, pk=None):
        material = self.get_object()
        
        if material.categoria != 'ENTORPECENTE':
            return Response({"error": "Apenas drogas passam por pesagem/cofre."}, status=400)
            
        peso_real = request.data.get('peso_real')
        localizacao = request.data.get('localizacao_no_cofre', 'Cofre Principal')
        
        material.peso_real = peso_real
        material.localizacao_no_cofre = localizacao
        material.status = 'ARMAZENADO'
        material.save(update_fields=['peso_real', 'localizacao_no_cofre', 'status'])
        
        RegistroHistorico.objects.create(
            material=material,
            criado_por=get_auth_user(request),
            status_na_epoca='ARMAZENADO',
            observacao=f"CONFERÊNCIA FÍSICA: Pesagem Real {material.peso_formatado()}. Local: {localizacao}"
        )
        return Response({"message": "Material armazenado no cofre com sucesso!"})

    @action(detail=True, methods=['post'])
    def autorizar_incineracao(self, request, pk=None):
        material = self.get_object()
        
        if material.status != 'ARMAZENADO':
            return Response({"error": "O material precisa estar no cofre (ARMAZENADO)."}, status=400)
            
        material.status = 'AUTORIZADO'
        material.save(update_fields=['status'])
        
        RegistroHistorico.objects.create(
            material=material,
            criado_por=get_auth_user(request),
            status_na_epoca='AUTORIZADO',
            observacao="CONFERÊNCIA PROJUDI: Despacho judicial verificado. Pronto para incineração."
        )
        return Response({"message": "Material autorizado para destruição!"})

    @action(detail=True, methods=['post'])
    def confirmar_entrega_forum(self, request, pk=None):
        material = self.get_object()
        user = get_auth_user(request)
        
        file = request.FILES.get('recibo')
        if not file:
            return Response({"error": "O upload do recibo assinado é obrigatório."}, status=400)
            
        material.recibo_forum = file
        material.status = 'ENTREGUE_AO_JUDICIARIO'
        material.save()
        
        RegistroHistorico.objects.create(
            material=material,
            criado_por=user,
            status_na_epoca='ENTREGUE_AO_JUDICIARIO',
            observacao=f"ENTREGA NO FÓRUM CONFIRMADA. Recibo anexado: {file.name}"
        )
        return Response({"message": "Entrega confirmada e recibo armazenado!"})

class LoteIncineracaoViewSet(viewsets.ModelViewSet):
    queryset = LoteIncineracao.objects.all().order_by('-data_criacao')
    serializer_class = LoteIncineracaoSerializer
    permission_classes = [AllowAny]

    def perform_create(self, serializer):
        serializer.save(criado_por=get_auth_user(self.request))

    @action(detail=False, methods=['post'])
    def gerar_lotes_automaticos(self, request):
        # Seleciona drogas que o Juiz já autorizou (AUTORIZADO)
        materiais = Material.objects.filter(status='AUTORIZADO', categoria='ENTORPECENTE').order_by('data_criacao')
        
        processos = {}
        for mat in materiais:
            proc = mat.noticiado.ocorrencia.processo or mat.noticiado.ocorrencia.bou
            if proc not in processos:
                processos[proc] = []
            processos[proc].append(mat)
            
        processos_list = list(processos.keys())
        if not processos_list:
            return Response({"error": "Nenhum material aguardando lote."}, status=400)
            
        lotes_gerados = 0
        from datetime import datetime
        
        # Agrupa de 20 em 20 processos independentemente do número de sacolas/materiais
        for i in range(0, len(processos_list), 20):
            bloco_processos = processos_list[i:i+20]
            
            # Cria um lote
            identificador = f"LOTE-{datetime.now().strftime('%Y%m%d%H%M%S')}-{lotes_gerados+1}"
            lote = LoteIncineracao.objects.create(
                identificador=identificador,
                status='ABERTO',
                criado_por=get_auth_user(request)
            )
            lotes_gerados += 1
            
            # Associa os materiais daquele bloco de processos ao Lote
            for proc in bloco_processos:
                for mat in processos[proc]:
                    mat.lote = lote
                    mat.status = 'AGUARDANDO_INCINERACAO'
                    mat.save()
                    RegistroHistorico.objects.create(
                        material=mat, criado_por=get_auth_user(request), 
                        status_na_epoca='AGUARDANDO_INCINERACAO', 
                        observacao=f"Atribuído ao Lote de Queima Automático {lote.identificador}."
                    )
                    
        return Response({"message": f"{lotes_gerados} lotes criados com sucesso cobrindo até {len(processos_list)} processos."})

    @action(detail=True, methods=['post'])
    def adicionar_material_avulso(self, request, pk=None):
        lote = self.get_object()
        if lote.status != 'ABERTO':
            return Response({"error": "Lote fechado/incinerado."}, status=400)
            
        material_id = request.data.get('material_id')
        try:
            mat = Material.objects.get(id=material_id)
            mat.lote = lote
            mat.status = 'AGUARDANDO_INCINERACAO'
            mat.save()
            return Response({"message": "Adicionado com sucesso!"})
        except Material.DoesNotExist:
            return Response({"error": "Item Inexistente."}, status=404)

    @action(detail=True, methods=['get'])
    def imprimir_capa(self, request, pk=None):
        lote = self.get_object()
        caminho = gerar_capa_lote_pdf(lote, get_auth_user(request))
        file_url = request.build_absolute_uri(settings.MEDIA_URL + caminho)
        return Response({'url': file_url})

    @action(detail=True, methods=['post'])
    def finalizar_lote(self, request, pk=None):
        lote = self.get_object()
        
        if lote.status == 'INCINERADO':
            return Response({"error": "Lote já foi finalizado."}, status=status.HTTP_400_BAD_REQUEST)
        
        # Verificar se há arquivo assinado
        termo_assinado = request.FILES.get('termo_assinado')
        if not termo_assinado:
            return Response({"error": "Anexe o Termo de Destruição assinado."}, status=status.HTTP_400_BAD_REQUEST)
        
        # Salvar arquivo
        from django.utils import timezone
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        filename = f"termo_assinado_{lote.identificador}_{timestamp}.pdf"
        pasta = os.path.join(settings.MEDIA_ROOT, 'lotes', 'assinados')
        os.makedirs(pasta, exist_ok=True)
        caminho = os.path.join(pasta, filename)
        
        with open(caminho, 'wb+') as dest:
            for chunk in termo_assinado.chunks():
                dest.write(chunk)
        
        # Marcar lote como incinerado
        lote.status = 'INCINERADO'
        lote.data_incineracao = timezone.now()
        lote.save()
        
        # Marcar todos os materiais do lote como INCINERADO
        for mat in lote.materiais.all():
            mat.status = 'INCINERADO'
            mat.save()
            RegistroHistorico.objects.create(
                material=mat,
                criado_por=get_auth_user(request),
                status_na_epoca='INCINERADO',
                observacao=f"Destruído via Termo Assinado - Lote {lote.identificador}"
            )
        
        return Response({
            "message": "Lote finalizado com sucesso!",
            "lote_status": lote.status,
            "itens_finalizados": lote.materiais.count()
        })

    @action(detail=False, methods=['post'])
    def finalizar_massa(self, request):
        """Finaliza vários lotes de uma vez com um único upload"""
        lote_ids = request.data.get('lote_ids', [])
        protocolo = request.data.get('protocolo', '')
        
        if not lote_ids:
            return Response({"error": "Nenhum lote selecionado."}, status=400)
            
        termo_assinado = request.FILES.get('termo_assinado')
        if not termo_assinado:
            return Response({"error": "Anexe o documento de incineração assinado."}, status=400)
            
        lotes = LoteIncineracao.objects.filter(id__in=lote_ids, status='ABERTO')
        if not lotes.exists():
            return Response({"error": "Nenhum lote aberto encontrado."}, status=400)
            
        # Salvar arquivo uma única vez
        from django.utils import timezone
        import os
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        filename = f"remessa_incineracao_{timestamp}.pdf"
        pasta = os.path.join(settings.MEDIA_ROOT, 'lotes', 'assinados')
        os.makedirs(pasta, exist_ok=True)
        caminho = os.path.join(pasta, filename)
        
        with open(caminho, 'wb+') as dest:
            for chunk in termo_assinado.chunks():
                dest.write(chunk)
                
        count_materiais = 0
        for lote in lotes:
            lote.status = 'INCINERADO'
            lote.data_incineracao = timezone.now()
            lote.eprotocolo_geral = protocolo # Salva o protocolo geral no lote
            lote.save()
            
            for mat in lote.materiais.all():
                mat.status = 'INCINERADO'
                mat.save()
                RegistroHistorico.objects.create(
                    material=mat, criado_por=get_auth_user(request),
                    status_na_epoca='INCINERADO',
                    observacao=f"INCINERAÇÃO: Protocolo {protocolo}. Lote {lote.identificador}."
                )
                count_materiais += 1
                
        return Response({
            "message": f"{lotes.count()} lotes finalizados com sucesso!",
            "materiais_baixados": count_materiais
        })

    @action(detail=False, methods=['get'])
    def imprimir_capas_em_massa(self, request):
        """Gera um único PDF com todas as capas dos lotes selecionados"""
        lote_ids = request.query_params.getlist('lote_ids')
        if not lote_ids:
            return Response({"error": "Selecione os lotes."}, status=400)
            
        from .documentos_services import gerar_capas_lote_coletivas
        pdf_content = gerar_capas_lote_coletivas(lote_ids)
        
        from django.http import HttpResponse
        response = HttpResponse(pdf_content, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="capas_lotes_selecionados.pdf"'
        return response

    @action(detail=False, methods=['get'])
    def imprimir_certidao_por_vara(self, request):
        """Gera uma certidão de incineração agrupando todos os lotes ABERTOS por Vara"""
        lotes_abertos = LoteIncineracao.objects.filter(status='ABERTO')
        if not lotes_abertos.exists():
            return Response({"error": "Nenhum lote aberto encontrado."}, status=400)
            
        from .documentos_services import gerar_certidao_incineracao_coletiva
        filename = gerar_certidao_incineracao_coletiva(lotes_abertos, get_auth_user(request))
        
        file_url = request.build_absolute_uri(settings.MEDIA_URL + f"incineracao/{filename}")
        return Response({'url': file_url})

    @action(detail=True, methods=['post'])
    def remover_material(self, request, pk=None):
        """Remove um material do lote e volta para status AUTORIZADO (avulso)"""
        material_id = request.data.get('material_id')
        try:
            mat = Material.objects.get(id=material_id, lote_id=pk)
            mat.lote = None
            mat.status = 'AUTORIZADO'
            mat.save()
            return Response({"message": "Material removido do lote."})
        except Material.DoesNotExist:
            return Response({"error": "Material não encontrado neste lote."}, status=404)


class NaturezaPenalViewSet(viewsets.ModelViewSet):
    queryset = NaturezaPenal.objects.all()
    serializer_class = NaturezaPenalSerializer
    permission_classes = [AllowAny]
    search_fields = ['nome']

    @action(detail=False, methods=['get'])
    def autocomplete(self, request):
        termo = request.query_params.get('q', '')
        if termo:
            naturezas = self.queryset.filter(nome__icontains=termo)
        else:
            naturezas = self.queryset.all()[:20]
        serializer = self.get_serializer(naturezas, many=True)
        return Response(serializer.data)


class EquipePMViewSet(viewsets.ModelViewSet):
    queryset = EquipePM.objects.all()
    serializer_class = EquipePMSerializer
    permission_classes = [AllowAny]

    @action(detail=False, methods=['get'])
    def autocomplete(self, request):
        termo = request.query_params.get('q', '')
        if termo:
            equipe = self.queryset.filter(nome__icontains=termo)
        else:
            equipe = self.queryset.all()[:20]
        serializer = self.get_serializer(equipe, many=True)
        return Response(serializer.data)
